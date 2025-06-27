
# Exports Requests and Responses from Burp.html reports.
# Exports in to Word format.
# Updated to place the words Request and Response in to table boxes
# To do - add header formating



import sys
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def extract_url_from_request(text):
    lines = text.strip().splitlines()
    if lines:
        parts = lines[0].split()
        if len(parts) > 1:
            return parts[1]
    return "<unknown>"

def process_html_to_docx(input_html, output_docx):
    with open(input_html, "r", encoding="utf-8") as file:
        soup = BeautifulSoup(file, "html.parser")

    doc = Document()
    doc.add_heading('Appendix', level=1)
    doc.add_paragraph('Request & Response')

    titles = soup.find_all("div", class_="BODH0")
    rr_blocks = soup.find_all("div", class_="rr_div")

    for i in range(0, len(rr_blocks) - 1, 2):
        title_text = titles[i // 2].get_text(strip=True) if i // 2 < len(titles) else f"Issue {i // 2 + 1}"
        request_div = rr_blocks[i]
        response_div = rr_blocks[i + 1]

        request_text = request_div.get_text().strip()
        affected_url = extract_url_from_request(request_text)

        doc.add_paragraph(f"{title_text} {affected_url}", style='Normal')

        # Build 4-row table: Request label, Request, Response label, Response
        table = doc.add_table(rows=4, cols=1)
        table.style = 'Table Grid'

        # Row 1: "HTTP Request"
        cell1 = table.cell(0, 0)
        para1 = cell1.paragraphs[0]
        run1 = para1.add_run("HTTP Request")
        run1.bold = True
        run1.font.name = 'Arial'
        run1.font.size = Pt(10)
        para1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Row 2: Request body
        cell2 = table.cell(1, 0)
        para2 = cell2.paragraphs[0]
        for element in request_div.descendants:
            if isinstance(element, str):
                run = para2.add_run(element)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            elif element.name == "span" and "HIGHLIGHT" in element.get("class", []):
                run = para2.add_run(element.get_text())
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 0, 0)

        # Row 3: "HTTP Response"
        cell3 = table.cell(2, 0)
        para3 = cell3.paragraphs[0]
        run3 = para3.add_run("HTTP Response")
        run3.bold = True
        run3.font.name = 'Arial'
        run3.font.size = Pt(10)
        para3.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Row 4: Response body
        cell4 = table.cell(3, 0)
        para4 = cell4.paragraphs[0]
        for element in response_div.descendants:
            if isinstance(element, str):
                run = para4.add_run(element)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            elif element.name == "span" and "HIGHLIGHT" in element.get("class", []):
                run = para4.add_run(element.get_text())
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 0, 0)

    doc.save(output_docx)

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python burphtml_to_docx.py <input_html_file> <output_docx_file>")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]
    process_html_to_docx(input_file, output_file)
